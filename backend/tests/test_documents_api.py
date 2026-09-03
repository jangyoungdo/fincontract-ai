import json
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfReader, PdfWriter

from app.config import get_settings
from app.main import app
from app.models import AnalysisRecord, AuditEvent, DocumentRecord, get_session_factory

SAMPLE = "제1조 은행은 필요하다고 인정하는 경우 서비스 내용을 일방적으로 변경할 수 있다."


def test_analysis_record_is_committed_before_queue_publication(monkeypatch) -> None:
    """Prevent an idle worker from consuming an ID hidden by the request transaction."""
    published: list[str] = []

    class CommitCheckingRedis:
        def set(self, key: str, value: str, ex: int) -> None:
            assert key and value and ex > 0

        def rpush(self, key: str, analysis_id: str) -> None:
            with get_session_factory()() as session:
                assert session.get(AnalysisRecord, analysis_id) is not None
            published.append(analysis_id)

    settings = get_settings().model_copy(update={"use_redis": True})
    monkeypatch.setattr("app.api.documents.get_settings", lambda: settings)
    monkeypatch.setattr("app.api.documents.get_redis", CommitCheckingRedis)
    with TestClient(app) as client:
        uploaded = client.post(
            "/api/v1/documents",
            files={"file": ("queue-order.txt", SAMPLE.encode(), "text/plain")},
        )
        response = client.post(
            f"/api/v1/documents/{uploaded.json()['id']}/analyses",
            json={"experiment_arm": "A"},
        )
    assert response.status_code == 202
    assert published == [response.json()["id"]]


def test_txt_upload_analysis_report_and_delete() -> None:
    with TestClient(app) as client:
        uploaded = client.post(
            "/api/v1/documents",
            files={"file": ("terms.txt", SAMPLE.encode(), "text/plain")},
        )
        assert uploaded.status_code == 201, uploaded.text
        document_id = uploaded.json()["id"]
        with get_session_factory()() as session:
            stored = session.get(DocumentRecord, document_id)
            assert stored is not None
            encrypted = Path(stored.storage_path).read_bytes()
            assert SAMPLE.encode() not in encrypted
            assert stored.storage_path.endswith(".txt.enc")

        analyzed = client.post(
            f"/api/v1/documents/{document_id}/analyses", json={"experiment_arm": "D"}
        )
        assert analyzed.status_code == 201, analyzed.text
        body = analyzed.json()
        assert body["status"] == "completed"
        assert body["experiment_arm"] == "full"
        assert body["result"]["experiment"]["mode"] == "full"
        assert len(body["result"]["findings"]) == 1

        report = client.get(f"/api/v1/analyses/{body['id']}/report")
        assert report.status_code == 200
        assert "법률 판단이 아닌" in report.json()["disclaimer"]

        pdf_report = client.get(f"/api/v1/analyses/{body['id']}/report.pdf")
        assert pdf_report.status_code == 200
        assert pdf_report.headers["content-type"] == "application/pdf"
        reader = PdfReader(BytesIO(pdf_report.content))
        assert len(reader.pages) >= 1
        report_text = "\n".join(page.extract_text() or "" for page in reader.pages)
        assert "왜 문제 후보인가" in report_text
        assert "예상되는 고객 영향" in report_text
        assert "탐지된 위험 구조" in report_text
        assert "수정 방향" in report_text
        assert "검토용 예시 문안" in report_text
        assert "법적 근거 후보" in report_text
        assert "검증 상세" in report_text

        deleted = client.delete(f"/api/v1/documents/{document_id}")
        assert deleted.status_code == 200
        assert deleted.json()["status"] == "deleted"
        with get_session_factory()() as session:
            event_types = {
                event.event_type
                for event in session.query(AuditEvent).filter(AuditEvent.document_id == document_id)
            }
        assert {"document_uploaded", "analysis_created", "analysis_completed", "report_generated", "document_deleted"}.issubset(event_types)


def test_public_result_strips_legacy_full_text_and_preview_is_owned_and_deleted(tmp_path) -> None:
    settings = get_settings()
    document_id = "preview-owner-document"
    analysis_id = "preview-owner-analysis"
    preview_id = "a" * 24
    encrypted_path = tmp_path / "preview-owner.txt.enc"
    encrypted_path.write_bytes(b"encrypted-placeholder")
    preview_file = settings.report_dir / "previews" / analysis_id / f"{preview_id}.png"
    preview_file.parent.mkdir(parents=True, exist_ok=True)
    preview_file.write_bytes(b"\x89PNG\r\n\x1a\nmasked-pixels")
    now = datetime.now(timezone.utc)
    result = {
        "document": {"masked_text": "[NAME_1]의 문서 전문", "pii_replacement_count": 1},
        "warnings": [
            "LLM_SCHEMA_INVALID",
            "OPENAI_CONTEXT_REVIEW_FAILED",
            "mock 에이전트 결과는 실제 LLM 품질 평가에 사용할 수 없습니다.",
            "법적 근거는 검증 전 후보이며 원문·시행일 확인이 필요합니다.",
            "이 결과는 법률 판단이 아닌 검토 보조 자료입니다.",
        ],
        "findings": [
            {
                "source": {
                    "masked_text": "[NAME_1]의 조항 조각",
                    "preview_status": "available",
                    "preview_ids": [preview_id],
                }
            }
        ],
        "candidate_findings": [],
    }
    with get_session_factory()() as session:
        session.add(
            DocumentRecord(
                id=document_id,
                original_filename="preview-owner.txt",
                mime_type="text/plain",
                sha256="1" * 64,
                status="ready",
                storage_path=str(encrypted_path),
                uploaded_at=now,
                expires_at=now + timedelta(hours=1),
            )
        )
        session.add(
            AnalysisRecord(
                id=analysis_id,
                document_id=document_id,
                status="completed",
                disposition="ready_for_review",
                experiment_arm="full",
                result_json=json.dumps(result, ensure_ascii=False),
            )
        )
        session.commit()

    with TestClient(app) as client:
        public_result = client.get(f"/api/v1/analyses/{analysis_id}")
        preview = client.get(
            f"/api/v1/analyses/{analysis_id}/source-previews/{preview_id}.png"
        )
        unowned = client.get(
            f"/api/v1/analyses/{analysis_id}/source-previews/{'b' * 24}.png"
        )
        deleted = client.delete(f"/api/v1/documents/{document_id}")
        expired_preview = client.get(
            f"/api/v1/analyses/{analysis_id}/source-previews/{preview_id}.png"
        )

    assert "masked_text" not in public_result.json()["result"]["document"]
    assert public_result.json()["result"]["warnings"] == ["OPENAI_CONTEXT_REVIEW_FAILED"]
    assert preview.status_code == 200
    assert preview.headers["content-type"] == "image/png"
    assert preview.headers["cache-control"] == "no-store"
    assert unowned.status_code == 404
    assert deleted.status_code == 200
    assert expired_preview.status_code == 404
    assert not preview_file.exists()


def test_rejects_spoofed_pdf() -> None:
    with TestClient(app) as client:
        response = client.post(
            "/api/v1/documents",
            files={"file": ("fake.pdf", b"not a pdf", "application/pdf")},
        )
    assert response.status_code == 400
    assert "시그니처" in response.json()["detail"]


def test_real_docx_upload_extracts_and_analyzes_text() -> None:
    stream = BytesIO()
    document = Document()
    document.add_paragraph(SAMPLE)
    document.save(stream)
    with TestClient(app) as client:
        uploaded = client.post(
            "/api/v1/documents",
            files={"file": ("terms.docx", stream.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert uploaded.status_code == 201, uploaded.text
        analyzed = client.post(f"/api/v1/documents/{uploaded.json()['id']}/analyses", json={"experiment_arm": "A"})
    assert analyzed.status_code == 201, analyzed.text
    assert len(analyzed.json()["result"]["findings"]) == 1


def test_docx_table_text_is_extracted_in_document_order() -> None:
    stream = BytesIO()
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "변경 조항"
    table.cell(0, 1).text = SAMPLE
    document.save(stream)
    with TestClient(app) as client:
        uploaded = client.post(
            "/api/v1/documents",
            files={
                "file": (
                    "table-terms.docx",
                    stream.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert uploaded.status_code == 201, uploaded.text
        analyzed = client.post(
            f"/api/v1/documents/{uploaded.json()['id']}/analyses",
            json={"experiment_arm": "A"},
        )
    assert analyzed.status_code == 201, analyzed.text
    assert len(analyzed.json()["result"]["findings"]) == 1


def test_real_pdf_upload_extracts_and_analyzes_text() -> None:
    # A compact, valid one-page PDF fixture with a literal text stream.
    pdf = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj
3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj
4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
5 0 obj<</Length 95>>stream
BT /F1 12 Tf 72 720 Td (Bank may unilaterally change service terms when necessary.) Tj ET
endstream endobj
xref
0 6
0000000000 65535 f\x20
0000000009 00000 n\x20
0000000058 00000 n\x20
0000000115 00000 n\x20
0000000251 00000 n\x20
0000000321 00000 n\x20
trailer<</Size 6/Root 1 0 R>>
startxref
468
%%EOF"""
    with TestClient(app) as client:
        uploaded = client.post("/api/v1/documents", files={"file": ("terms.pdf", pdf, "application/pdf")})
        assert uploaded.status_code == 201, uploaded.text
        # This English fixture confirms actual PDF extraction; it need not trigger Korean rules.
        analyzed = client.post(f"/api/v1/documents/{uploaded.json()['id']}/analyses", json={"experiment_arm": "A"})
    assert analyzed.status_code == 201, analyzed.text
    assert analyzed.json()["result"]["clause_count"] == 1


def test_scanned_pdf_fails_closed_with_ocr_required() -> None:
    stream = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.write(stream)
    with TestClient(app) as client:
        response = client.post(
            "/api/v1/documents",
            files={"file": ("scan.pdf", stream.getvalue(), "application/pdf")},
        )
    assert response.status_code == 400
    assert response.json()["detail"].startswith("OCR_REQUIRED:")


def test_bank_comparison_fails_closed_without_a_verified_dataset() -> None:
    with TestClient(app) as client:
        response = client.get("/api/v1/bank-comparisons")
    assert response.status_code == 503
    assert "비교 데이터" in response.json()["detail"]
