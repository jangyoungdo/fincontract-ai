from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

from app.config import get_settings


@dataclass(frozen=True)
class ExtractedDocument:
    """Page-aware extraction result; ``text`` remains the compatibility surface."""
    pages: tuple[str, ...]
    text: str
    extension: str


def _normalized_margin_line(line: str) -> str:
    return re.sub(r"\d+", "#", re.sub(r"\s+", " ", line.strip())).casefold()


def _remove_repeated_margins(pages: list[str]) -> list[str]:
    """Remove only lines repeated in the first/last three lines of multiple pages."""
    occurrences: dict[str, set[int]] = {}
    page_lines = [page.splitlines() for page in pages]
    for page_index, lines in enumerate(page_lines):
        margin_indexes = set(range(min(5, len(lines))))
        margin_indexes.update(range(max(0, len(lines) - 5), len(lines)))
        for index in margin_indexes:
            normalized = _normalized_margin_line(lines[index])
            if normalized:
                occurrences.setdefault(normalized, set()).add(page_index)
    repeated = {line for line, indexes in occurrences.items() if len(indexes) >= 2}
    cleaned = []
    for lines in page_lines:
        margin_indexes = set(range(min(5, len(lines))))
        margin_indexes.update(range(max(0, len(lines) - 5), len(lines)))
        kept = [line for index, line in enumerate(lines) if index not in margin_indexes or _normalized_margin_line(line) not in repeated]
        cleaned.append("\n".join(kept).strip())
    return cleaned


def _is_usable_ocr_text(text: str, minimum_characters: int, minimum_alnum_ratio: float) -> bool:
    """Reject empty or mostly-symbol OCR output before it reaches analysis."""
    compact = "".join(character for character in text if not character.isspace())
    if len(compact) < minimum_characters:
        return False
    alnum_count = sum(character.isalnum() for character in compact)
    return alnum_count / len(compact) >= minimum_alnum_ratio


def _ocr_pdf_pages(data: bytes, page_indexes: list[int], settings: Any) -> dict[int, str]:
    """Render selected PDF pages in memory and run bounded local Tesseract OCR."""
    try:
        import pypdfium2 as pdfium
        import pytesseract
    except ImportError as exc:
        raise ValueError("OCR_UNAVAILABLE: 로컬 OCR 구성요소를 불러올 수 없습니다.") from exc

    document = pdfium.PdfDocument(data)
    extracted: dict[int, str] = {}
    scale = settings.ocr_dpi / 72
    try:
        for page_index in page_indexes:
            page = document[page_index]
            try:
                width, height = page.get_size()
                pixel_count = int(width * scale) * int(height * scale)
                if pixel_count > settings.ocr_max_pixels_per_page:
                    raise ValueError("OCR_PIXEL_LIMIT: OCR 페이지 픽셀 제한을 초과했습니다.")
                bitmap = page.render(scale=scale)
                try:
                    with bitmap.to_pil() as image:
                        try:
                            text = pytesseract.image_to_string(
                                image,
                                lang=settings.ocr_languages,
                                config="--oem 1 --psm 6",
                                timeout=settings.ocr_timeout_seconds,
                            )
                        except pytesseract.pytesseract.TesseractNotFoundError as exc:
                            raise ValueError(
                                "OCR_UNAVAILABLE: Tesseract 실행 파일을 찾을 수 없습니다."
                            ) from exc
                        except pytesseract.pytesseract.TesseractError as exc:
                            raise ValueError(
                                "OCR_UNAVAILABLE: Tesseract 언어 모델을 사용할 수 없습니다."
                            ) from exc
                        except RuntimeError as exc:
                            raise ValueError(
                                "OCR_TIMEOUT: OCR 페이지 처리 시간을 초과했습니다."
                            ) from exc
                finally:
                    bitmap.close()
            finally:
                page.close()
            cleaned = text.replace("\x00", "").strip()
            if not _is_usable_ocr_text(
                cleaned,
                settings.ocr_min_characters_per_page,
                settings.ocr_min_alnum_ratio,
            ):
                raise ValueError("OCR_LOW_CONFIDENCE: OCR 텍스트 품질 기준을 통과하지 못했습니다.")
            extracted[page_index] = cleaned
    finally:
        document.close()
    return extracted


def _extract_pdf_document(data: bytes) -> ExtractedDocument:
    """Combine native page text with local OCR only for pages that need it."""
    settings = get_settings()
    reader = PdfReader(BytesIO(data))
    if reader.is_encrypted:
        raise ValueError("PDF_ENCRYPTED: 암호화된 PDF는 처리할 수 없습니다.")
    if len(reader.pages) > settings.pdf_max_pages:
        raise ValueError("PDF_PAGE_LIMIT: PDF 페이지 제한을 초과했습니다.")

    page_texts = [(page.extract_text() or "").strip() for page in reader.pages]
    native_text_exists = any(page_texts)
    ocr_indexes: list[int] = []
    for index, (page, text) in enumerate(zip(reader.pages, page_texts, strict=True)):
        if text:
            continue
        # Blank pages inside an otherwise readable PDF are harmless. Image-only
        # pages and wholly non-text PDFs require OCR so content is never dropped.
        has_images = bool(list(page.images))
        if has_images or not native_text_exists:
            ocr_indexes.append(index)

    if ocr_indexes:
        if not settings.ocr_enabled:
            raise ValueError("OCR_REQUIRED: 스캔 PDF에는 로컬 OCR이 필요합니다.")
        for index, text in _ocr_pdf_pages(data, ocr_indexes, settings).items():
            page_texts[index] = text
    page_texts = _remove_repeated_margins(page_texts)
    text = "\n".join(text for text in page_texts if text)
    return ExtractedDocument(tuple(page_texts), text, ".pdf")


def _extract_pdf_text(data: bytes) -> str:
    return _extract_pdf_document(data).text


def extract_document(
    data: bytes, extension: str, max_characters: int = 200_000
) -> ExtractedDocument:
    """Extract a bounded page-aware document while preserving the string contract."""
    if extension == ".txt":
        text = data.decode("utf-8")
        pages = (text,)
    elif extension == ".pdf":
        extracted = _extract_pdf_document(data)
        text = extracted.text
        pages = extracted.pages
    elif extension == ".docx":
        document = Document(BytesIO(data))
        blocks = []
        for block in document.iter_inner_content():
            if isinstance(block, Paragraph) and block.text.strip():
                blocks.append(block.text)
            elif isinstance(block, Table):
                for row in block.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        blocks.append(" | ".join(cells))
        text = "\n".join(blocks)
        pages = (text,)
    else:
        raise ValueError("지원하지 않는 추출 형식입니다.")

    text = text.replace("\x00", "").strip()
    if not text:
        raise ValueError("문서에서 텍스트를 추출하지 못했습니다.")
    if len(text) > max_characters:
        raise ValueError("추출 문자 수 제한을 초과했습니다.")
    cleaned_pages = tuple(page.replace("\x00", "").strip() for page in pages)
    return ExtractedDocument(cleaned_pages, text, extension)


def extract_text(data: bytes, extension: str, max_characters: int = 200_000) -> str:
    """Compatibility wrapper for callers that still require one flattened string."""
    return extract_document(data, extension, max_characters).text
